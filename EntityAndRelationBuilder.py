from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LAParams, LTTextBox, LTTextLine
from DBHandler import DBHandler
import re
import docx
import spacy
import Translator
from nltk.corpus import stopwords
stop_words = set(stopwords.words('english'))
print(stop_words)
from pycorenlp import *
import time
from RelationTriple import RelationTriple
#output_dir = "D:\\Crymzee\\Steve\\Model\\attomusmodel"
#nlp = spacy.load('en_core_web_sm')
#nlp = spacy.load(output_dir)
import json
from Entity import Entity
from collections import Counter
import re
import traceback





class EntityAndRelationBuilder(object):

    def __init__(self, lang, filename, projectName,document_url ,fname,title):
        with open('configurations.json') as f:
            data = json.load(f)
        self.lang = lang
        self.filename = filename
        self.document_url = document_url
        self.projectName = projectName
        self.credentials = data["GOOGLE_API_KEY"]
        self.StanfordNLPPath = data["NLP_API_KEY"]
        self.relationExtractor = StanfordCoreNLP(self.StanfordNLPPath)
        self.DATABASEIP = data["DATABASEIP"]
        self.DB_USER = data["DB_USER"]
        self.DB_PASSWORD = data["DB_PASSWORD"]
        self.DATABASE = data["DATABASE"]
        self.fname = fname
        self.entitiesall = []
        self.uniqueentities = []
        self.title = title
        self.personEntities = []
        self.orgEntities = []
        self.output_dir = data["MODEL"]
        self.nlp = spacy.load('en_core_web_sm')
        if(self.output_dir not in "" or self.output_dir !=None):
            self.nlp = spacy.load(self.output_dir)



    def __del__(self):
        print("destructor")

    def getEntityPDFJson(self):

        try:
            fp = open(self.filename, 'rb')
            parser = PDFParser(fp)
            doc = PDFDocument()
            parser.set_document(doc)
            doc.set_parser(parser)
            doc.initialize('')
            rsrcmgr = PDFResourceManager()
            laparams = LAParams()
            laparams.word_margin = 1.0
            device = PDFPageAggregator(rsrcmgr, laparams=laparams)
            interpreter = PDFPageInterpreter(rsrcmgr, device)
            extracted_text = ''
            isEnglish = True
            relationsList = []
            relation = RelationTriple("", "", "", self.document_url, self.fname, self.projectName)
            relationsList.append(relation)
            titleExtracted = False
            for page in doc.get_pages():
                interpreter.process_page(page)
                layout = device.get_result()
                fullText =""
                extracted_text =""
                for lt_obj in layout:
                    if isinstance(lt_obj, LTTextBox) or isinstance(lt_obj, LTTextLine):
                        extracted_text += lt_obj.get_text()

                if(not titleExtracted):
                    n1 = extracted_text.replace("\t", " ")
                    n2 = n1.replace("\r", " ")
                    finaltext = n2.replace("\u00a0", " ")
                    lines = finaltext.split("\n")
                    for line in lines :
                        if (not titleExtracted):
                            if(re.search('[a-zA-Z]', line)):
                                print("Line:" + line)
                                titleExtracted = True
                                self.title = line
                                print("Title Extracted is :"+self.title)


                n1 = extracted_text.replace("\t", " ")
                n2 = n1.replace("\r", " ")
                n3 = n2.replace("\n", " ")
                finaltext = n3.replace("\u00a0", " ")
                fullText = fullText + finaltext

                print("Text Extracted")
                print(fullText)
                print("Language is :"+self.lang)
                if (self.lang not in "eng"):
                    print("Going to translate")
                    fullText = Translator.translate_(fullText, self.credentials)
                    time.sleep(100)
                paras = fullText.split(".")
                print("Number of Para:" + str(len(paras)))
                entities = self.extractentities(fullText)

                for para in paras:
                    relations = self.extractrelation(para)
                    print(relations)
                    for r in relations:
                        relation = RelationTriple(r.getSubject_(), r.getObject_(), r.getRelation_(),
                                                  self.document_url, self.fname, self.projectName)

                        relationsList.append(relation)
                        del relation

            print("dhgjdhgjdhgdjgh gsdhguihg guityweruity")
            res = self.sourceRelations()
            response = res.split("(");
            relationsList[0].subject_ = response[0]
            relationsList[0].object_ = response[1]
            relationsList[0].relation_ = self.title
            for temp in self.uniqueentities:
                print("Entity:"+temp)
                relation = RelationTriple(temp, "", "", self.document_url, self.fname,
                                          self.projectName)
                relationsList.append(relation)
                del relation

            for rel in relationsList:
                print(self.title)
                rel.source = response[0]
                rel.sourceSubject = self.title


            print("Relations Count:" + str(len(relationsList)))
            dbHandler_ = DBHandler(self.DATABASEIP, self.DB_USER, self.DB_PASSWORD, self.DATABASE)
            dbHandler_.insertRelation(relationsList)
            dbHandler_.updateFileStatus(self.fname, "Y")

        except Exception as e:
            print("Exception in Processing file"+self.fname+str(e))
            traceback.print_exc()

            dbHandler_.updateFileStatus(self.fname, "F")

        return True

    def getEntityDocxJson(self):
        print("docx file")

        doc = docx.Document(self.filename)
        relationsList = []
        relation = RelationTriple("", "", "", self.document_url, self.fname, self.projectName)
        relationsList.append(relation)
        for para in doc.paragraphs:
            if (self.lang not in "eng"):
                translatedText = Translator.translate_(para.text, self.credentials)
                para.text = translatedText
                time.sleep(100)
            entities = self.extractentities(para.text)
            relations = self.extractrelation(para.text)
            print(entities)
            print(relations)
            '''
            for entity in entities:
                relation = RelationTriple(entity.Text, "", "", self.document_url,self.fname,self.projectName)
                relationsList.append(relation)
                del relation
            '''
            for r in relations:
                relation = RelationTriple(r.getSubject_(), r.getObject_(), r.getRelation_(), self.document_url,self.fname,self.projectName)
                relationsList.append(relation)
                del relation
        res = self.sourceRelations()
        response = res.split("(");
        relationsList[0].subject_ = response[0]
        relationsList[0].object_ = response[1]
        relationsList[0].relation_ = self.title
        for temp in self.uniqueentities:
            print("Entity:" + temp)
            relation = RelationTriple(temp, "", "", self.document_url, self.fname,
                                      self.projectName)
            relationsList.append(relation)
            del relation

        for rel in relationsList:
            print(self.title)
            rel.source = response[0]
            rel.sourceSubject = self.title

        dbHandler_ = DBHandler(self.DATABASEIP, self.DB_USER, self.DB_PASSWORD, self.DATABASE)
        dbHandler_.insertRelation(relationsList)
        dbHandler_.updateFileStatus(self.fname, "Y")

        return True

    def getEntityTxtJson(self):
        print("Txt file")
        f = open(self.filename, 'rb')
        data = f.read().decode('utf8', 'ignore')
        paragraphs = data.split("\n\n")
        translatedPara = ""
        relationsList = []
        relation = RelationTriple("", "", "", self.document_url, self.fname, self.projectName)
        relationsList.append(relation)

        counter = 1
        print("No of Paragraphs")
        print(len(paragraphs))

        for para in paragraphs:

            print("Size of Paragraph")
            print(len(para))
            paras =[]
            if(len(para)>20000):
                #paras = splitCount(para,20000)
                paras = split(para)
            else:
                paras.append(para)

            print("Number of Para")
            print(len(paras))
            for p in paras:
                print("Size of Para")
                print(len(p))
                if (self.lang not in "eng"):
                    print("Going to call Google API for Translation")
                    translatedPara = Translator.translate_(p, self.credentials)
                    print(translatedPara)
                    print("Going to sleep for 1 min as google API does not allow call within 1 minute")
                    p = translatedPara
                    time.sleep(100)
                entities = self.extractentities(p)
                relations = self.extractrelation(p)

                for r in relations:
                    relation = RelationTriple(r.getSubject_(), r.getObject_(), r.getRelation_(), self.document_url,self.fname,self.projectName)
                    relationsList.append(relation)

                    del relation

        res = self.sourceRelations()
        response = res.split("(");
        relationsList[0].subject_ = response[0]
        relationsList[0].object_ = response[1]
        relationsList[0].relation_ = self.title
        for temp in self.uniqueentities:
            print("Entity:" + temp)
            relation = RelationTriple(temp, "", "", self.document_url, self.fname,
                                      self.projectName)
            relationsList.append(relation)
            del relation

        for rel in relationsList:
            print(self.title)
            rel.source = response[0]
            rel.sourceSubject = self.title

        dbHandler_ =DBHandler(self.DATABASEIP,self.DB_USER,self.DB_PASSWORD,self.DATABASE)
        dbHandler_.insertRelation(relationsList)
        dbHandler_.updateFileStatus(self.fname,"Y")
        return True


    def isSearchablePDF(self):
        searchable = True
        entityList_ = []
        print("PDF File")
        fp = open(self.filename, 'rb')
        print("PDF Filejhhfyf")

        parser = PDFParser(fp)
        doc = PDFDocument()
        parser.set_document(doc)
        doc.set_parser(parser)
        doc.initialize('')
        rsrcmgr = PDFResourceManager()
        laparams = LAParams()
        laparams.char_margin = 1.0
        laparams.word_margin = 1.0
        device = PDFPageAggregator(rsrcmgr, laparams=laparams)
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        extracted_text = ''
        counter = 1
        try:
            for page in doc.get_pages():
                interpreter.process_page(page)
                layout = device.get_result()
                for lt_obj in layout:
                    if isinstance(lt_obj, LTTextBox) or isinstance(lt_obj, LTTextLine):
                        extracted_text = lt_obj.get_text()
                        entities = self.extractentities(extracted_text)
                        for entity in entities:
                            entityList_.append(entity)
            if entityList_ == []:
                searchable = False

        except Exception as e:
            print("File is not parsaing")

        return searchable;


    def isAlreadyThere(self, entityList, Text):
        #Text = re.sub('[^A-Za-z0-9]+', '', Text)
        if entityList != []:
            for entity in entityList:
                #entity = re.sub('[^A-Za-z0-9]+', '', entity)
                if (Text.lower() == entity.lower() or Text.lower().__contains__(
                        entity.lower()) or entity.lower().__contains__(Text.lower())):
                    return True
            return False
        else:
            return False



    def extractentities(self,text):
        doc = self.nlp(text)
        entities = []
        for X in doc.ents:
            if X.text != ('\n')  and X.label_  in ('PERSON', 'ORG', 'LOC', 'DATE','MONEY','GPE'):
                print(X.text,X.label_)
                entity = X.text
                entity = re.sub('[^A-Za-z0-9]+', '', entity)
                n1 = entity.replace("\t", " ")
                n2 = n1.replace("\r", " ")
                n3 = n2.replace("\n", " ")
                finaltext = n3.replace("\u00a0", " ")
                finaltext_ = finaltext.replace("\xa0", " ")
                finaltext_.lstrip(" ")
                finaltext_.rstrip(" ")
                entity = finaltext_

                if(entity.lower() in stop_words or entity in " " or entity in "_" or entity.isdigit() or not self.verify(entity)):
                    print("Entity Skipped")
                else:
                    entity_ = Entity(entity,X.label_)
                    entities.append(entity_)
                    self.entitiesall.append(X.text)
                    if (not self.isAlreadyThere(self.uniqueentities,entity)):
                        self.uniqueentities.append(entity)
                    if(X.label_ == 'PERSON'):
                        self.personEntities.append(entity)
                    elif(X.label_ == 'ORG'):
                        self.orgEntities.append(entity)
        return  entities

    def extractrelation(self,text):
        relationsList = []
        text.replace("\n","")
        for line in text.split("."):
            output = self.relationExtractor.annotate(line, properties={
                "annotators": "tokenize,ssplit,pos,depparse,natlog,openie",
                "outputFormat": "json",
                "openie.triple.strict": "true",
                "openie.max_entailments_per_clause": "1"})
            if (output != None):
                if (len(output["sentences"]) > 0):
                    result = [output["sentences"][0]["openie"] for item in output]
                    # print(len(result))
                    for i in result:
                        for rel in i:
                            relationSent = rel['subject'], rel['relation'], rel['object']
                            print(relationSent)
                            relation = RelationTriple(rel['subject'], rel['object'], rel['relation'],
                                                      self.document_url,self.fname,self.projectName)
                            relationsList.append(relation)
                            del relation

        return relationsList

    def updaterelations(self,entities):
        updatedrelations = []
        person = self.getPersonEntity(entities)
        organization = self.getOrgEntity(entities)
        if person != None:
            for entity in entities:
                if entity.Label in "LOC" or entity.Label in "GPE":
                    relation = RelationTriple(person.Text, entity.Text, "works at",
                                              self.document_url, self.fname, self.projectName)
                    updatedrelations.append(relation)
        elif organization != None:
            for entity in entities:
                if entity.Label in "LOC" or entity.Label in "GPE":
                    relation = RelationTriple(organization.Text, entity.Text, "is located ",
                                              self.document_url, self.fname, self.projectName)

                    updatedrelations.append(relation)


        return updatedrelations

    def getPersonEntity(self,entities):
        for entity in entities:
            if entity.Label =="PERSON":
                return entity
        return None

    def getOrgEntity(self,entities):
        for entity in entities:
            if entity.Label =="ORG":
                return entity
        return None


    def sourceRelations(self):
        list = Counter(self.personEntities)
        source = list.most_common(2)
        subject = ""
        object = ""
        res = subject+"("+object
        try:
            subject = source[0][0]
            object = source[1][0]
            print("Source Subject:"+subject)
            print("Source Object:" + object)
            res = subject + "(" + object
        except Exception as e:
            print(str(e))

        return res;

    def verify(self,entity):
        valid = True
        entitytoken = entity.split(" ")
        count = 0
        stpowordCount = 0
        for t in entitytoken:
            if t in stop_words:
                stpowordCount = stpowordCount+1
            else:
                count = count+1

        if(stpowordCount > count):
            valid = False
        print("Entity:"+entity+" is"+str(valid))
        return valid



def splitCount(s, count):
    length = len(s)
    i=0
    startIndex = 0
    endIndex = 0
    splittedStr = []
    while endIndex < length:
        if(endIndex+count > length):
            break
        endIndex = endIndex+count
        splittedStr.append(s[startIndex:endIndex])
        startIndex = endIndex+1
    if (startIndex < length):
        splittedStr.append(s[startIndex:(length-1)])

    return splittedStr

def split(s):
    lines = s.split(".")
    paras = []
    count = 0
    para = ""
    for line in lines:
        if(count % 10 ==0):
            paras.append(para)
            para=""
        para.append(line)
        count=count+1
    paras.append(para)


    return paras







