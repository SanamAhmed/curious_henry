#from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfparser import *
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LAParams, LTTextBox, LTTextLine
import re
import docx
import spacy
import Translator
from pycorenlp import *
import pymysql
import time

from RelationTriple import RelationTriple

nlp = spacy.load('en_core_web_sm')
import json
class EntityExtractor(object):


    def __init__(self,lang, filename ,projectName,document_url,credentials,StanfordNLPPath, DATABASEIP,DB_USER,DB_PASSWORD,DATABASE,fname):
            self.lang = lang
            self.filename=filename
            self.nodesList = []
            self.edgesList = []
            self.listEntities = []
            self.document_url = document_url
            self.credentials = credentials
            self.StanfordNLPPath = StanfordNLPPath
            self.relationExtractor = StanfordCoreNLP(self.StanfordNLPPath)
            self.projectName = projectName

            self.DATABASEIP = DATABASEIP
            self.DB_USER = DB_USER
            self.DB_PASSWORD = DB_PASSWORD
            self.DATABASE = DATABASE
            self.fname = fname

    def __del__(self):
        print ("destructor")

    def getEntityPDFJson(self):
        #searchable = isSearchablePDF()
        counter=1

        print("PDF File")
        fp = open(self.filename, 'rb')
        parser = PDFParser(fp)
        doc = PDFDocument()
        parser.set_document(doc)
        doc.set_parser(parser)
        doc.initialize('')
        rsrcmgr = PDFResourceManager()
        laparams = LAParams()
        laparams.char_margin = 1.0
        laparams.word_margin = 1.0
        device = PDFPageAggregator(rsrcmgr, laparams=laparams)
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        extracted_text = ''
        isEnglish = True
        relationsList=[]
        uniqueEntities = []
        for page in doc.get_pages():
            interpreter.process_page(page)
            layout = device.get_result()
            for lt_obj in layout:

                if isinstance(lt_obj, LTTextBox) or isinstance(lt_obj, LTTextLine):

                    if(self.lang not in"eng"):
                        extracted_text =Translator.translate_(lt_obj.get_text(),self.credentials)
                    else:
                        extracted_text = lt_obj.get_text()

                    n1 = extracted_text.replace("\t", " ")
                    n2 = n1.replace("\r", "")
                    n3 = n2.replace("\n", "")
                    finaltext = n3.replace("\u00a0", "")
                    doc = nlp(finaltext)
                    paras = extracted_text.split("\n\n")
                    for p in paras:

                        for line in p.split("\n"):
                            output = self.relationExtractor.annotate(line, properties={
                                "annotators": "tokenize,ssplit,pos,depparse,natlog,openie",
                                "outputFormat": "json",
                                "openie.triple.strict": "true",
                                "openie.max_entailments_per_clause": "1"})
                            if (output != None):
                                if (len(output["sentences"]) > 0):
                                    result = [output["sentences"][0]["openie"] for item in output]
                                    # print(len(result))
                                    for i in result:
                                        for rel in i:
                                            relationSent = rel['subject'], rel['relation'], rel['object']
                                            print(relationSent)
                                            relation=RelationTriple(rel['subject'],rel['object'],rel['relation'],self.document_url)
                                            relationsList.append(relation)
                                            del relation

                    for X in doc.ents:
                        if X.text != ('\n') and X.label_ not in ('ORDINAL', 'CARDINAL','NORP','Non-­‐binding'):

                            if(self.isAlreadyThere(uniqueEntities,X.text)==False):
                                self.listEntities.append(X.text + ",")
                                uniqueEntities.append(X.text)




        print(self.listEntities)
        print("Relations Count:"+str(len(relationsList)))
        self.insertRelation(relationsList)

        return self.listEntities

    def getEntityDocxJson(self):
        print("docx file")
        doc = docx.Document(self.filename)
        relationsList = []
        translatedText =""
        uniqueEntities = []

        for para in doc.paragraphs:
            if(self.lang not in "eng"):
                translatedText = Translator.translate_(para.text,self.credentials)
                para.text=translatedText
                doc = nlp(translatedText)
                time.sleep(100)
            else:

                doc = nlp(para.text)
            # print([(X.text, X.label_) for X in doc.ents])
            # print(type([(X.text, X.label_) for X in doc.ents]))
            # list1 = [(X.text, X.label_) for X in doc.ents]
            for line in para.text.split("\n"):
                output = self.relationExtractor.annotate(line, properties={
                    "annotators": "tokenize,ssplit,pos,depparse,natlog,openie",
                    "outputFormat": "json",
                    "openie.triple.strict": "true",
                    "openie.max_entailments_per_clause": "1"})
                if (output != None):
                    if (len(output["sentences"]) > 0):
                        result = [output["sentences"][0]["openie"] for item in output]
                        # print(len(result))
                        for i in result:
                            for rel in i:
                                relationSent = rel['subject'], rel['relation'], rel['object']
                                rellation = RelationTriple(rel['subject'], rel['object'], rel['relation'],self.document_url)
                                relationsList.append(rellation)
                                del rellation
                                print(relationSent)

            for X in doc.ents:
                if X.text != ('\n') and X.label_ not in ('ORDINAL', 'CARDINAL','NORP','Non-­‐binding'):

                    if self.isAlreadyThere(uniqueEntities,X.text) == False:
                        uniqueEntities.append(X.text)
                        self.listEntities.append(X.text + " " + "" + X.label_ + ",")

                    else:
                        print("Repated Entity"+X.text+"Skipped")

        self.insertRelation(relationsList)

        return self.listEntities


    def getEntityTxtJson(self):
        print("Txt file")
        f = open(self.filename, 'rb')
        data = f.read().decode('utf8', 'ignore')
        paragraphs = data.split("\n\n")
        translatedPara=""
        relationsList = []
        counter=1
        uniqueEntities = []
        for para in paragraphs:
            if(self.lang not in "eng"):
                translatedPara = Translator.translate_(para,self.credentials)
                para = translatedPara

            doc = nlp(para)
            for line in para.split("\n"):
                output = self.relationExtractor.annotate(line, properties={
                    "annotators": "tokenize,ssplit,pos,depparse,natlog,openie",
                    "outputFormat": "json",
                    "openie.triple.strict": "true",
                    "openie.max_entailments_per_clause": "1"})
                if (output != None):
                    if (len(output["sentences"]) > 0):
                        result = [output["sentences"][0]["openie"] for item in output]
                        for i in result:
                            for rel in i:
                                relationSent = rel['subject'], rel['relation'], rel['object']
                                rellation = RelationTriple(rel['subject'], rel['object'], rel['relation'],self.document_url)
                                relationsList.append(rellation)
                                del rellation
                                print(relationSent)
            for X in doc.ents:
                if X.text != ('\n') and X.label_ not in ('ORDINAL', 'CARDINAL','NORP','Non-­‐binding'):
                    if self.isAlreadyThere(uniqueEntities,X.text) == False:
                        uniqueEntities.append(X.text)
                        self.listEntities.append(X.text + ",")

                    else:
                        print("Repated Entity"+X.text+"Skipped")


        self.insertRelation(relationsList)
        return self.listEntities




    def getNodesList(self):
        print(self.nodesList)
        return json.dumps(self.nodesList)

    def getEdgeList(self):
        print(self.edgesList)
        return json.dumps(self.edgesList)
    def getEntities(self):
        entitystr ="".join(self.listEntities)
        return entitystr

    def isSearchablePDF(self):
        searchable = True
        entityList_ = []
        print("PDF File")
        fp = open(self.filename, 'rb')
        parser = PDFParser(fp)
        doc = PDFDocument()
        parser.set_document(doc)
        doc.set_parser(parser)
        doc.initialize('')
        rsrcmgr = PDFResourceManager()
        laparams = LAParams()
        laparams.char_margin = 1.0
        laparams.word_margin = 1.0
        device = PDFPageAggregator(rsrcmgr, laparams=laparams)
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        extracted_text = ''
        counter =1

        for page in doc.get_pages():
            interpreter.process_page(page)
            layout = device.get_result()
            for lt_obj in layout:
                if isinstance(lt_obj, LTTextBox) or isinstance(lt_obj, LTTextLine):
                    extracted_text += lt_obj.get_text()

                    n1 = extracted_text.replace("\t", " ")
                    n2 = n1.replace("\r", "")
                    n3 = n2.replace("\n", "")
                    finaltext = n3.replace("\u00a0", "")

                    doc = nlp(finaltext)
                    # print([(X.text, X.label_) for X in doc.ents])
                    for X in doc.ents:
                        if X.text != ('\n') and X.label_ not in (
                        'ORDINAL', 'CARDINAL', 'NORP', 'Non-­‐binding'):
                            self.listEntities.append(X.text + ",")
                            entityList_.append((X.text, X.label_))

        if(entityList_==[]):
            searchable=False
        return searchable;


    def insertRelation(self,relationList):
        db = pymysql.connect(self.DATABASEIP, self.DB_USER, self.DB_PASSWORD,
                             self.DATABASE)
        cur = db.cursor()


        sql = 'INSERT INTO relations (ProjectName ,FileName,DocumentURL ,Subject , Object,Relation  ) VALUES (%s, %s,%s,%s,%s,%s)'

        for relation in relationList:

            args = (self.projectName, self.fname, self.document_url,relation.getSubject_(),relation.getObject_(),relation.getRelation_())
            try:

                # Execute the SQL command
                cur.execute(sql, args)
                # Commit your changes in the database
                db.commit()
                # db.close()


            except Exception as e:
                print(e)
        db.close()


    def isAlreadyThere(self,entityList ,Text):
        Text=re.sub('[^A-Za-z0-9]+', '', Text)
        if entityList !=[]:
            for entity in entityList:
                entity = re.sub('[^A-Za-z0-9]+', '', entity)
                if(Text.lower() == entity.lower() or Text.lower().__contains__(entity.lower()) or entity.lower().__contains__(Text.lower())):
                    return True
            return False
        else:
            return False








